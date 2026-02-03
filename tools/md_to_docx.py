import sys
import os
import re
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Paragraphs
        elif line:
            # Simple handling of bold/italic for paragraphs
            # Removes markdown symbols for now to keep it clean in docx
            clean_line = re.sub(r'\*\*|\*|`|\[.*?\]\(.*?\)', '', line)
            doc.add_paragraph(clean_line)
        else:
            doc.add_paragraph('')

    doc.save(docx_path)
    print(f"Converted {md_path} to {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <file1.md> <file2.md> ...")
        sys.exit(1)
    
    for md_file in sys.argv[1:]:
        docx_file = md_file.replace('.md', '.docx')
        md_to_docx(md_file, docx_file)
